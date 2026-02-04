import streamlit as st
import google.generativeai as genai
import fitz  # PyMuPDF
import docx
import json
import difflib
import re
import unicodedata
import time
from spellchecker import SpellChecker
from io import BytesIO

# ----------------- 1. VISUAL & CSS -----------------
st.set_page_config(page_title="Conferência MKT", page_icon="💊", layout="wide")

st.markdown("""
<style>
    [data-testid="stHeader"] { visibility: hidden; }

    .texto-box { 
        font-family: 'Segoe UI', sans-serif;
        font-size: 0.95rem;
        line-height: 1.7;
        color: #212529;
        background-color: #ffffff;
        padding: 25px;
        border-radius: 8px;
        border: 1px solid #ced4da;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        text-align: left;
    }

    .highlight-yellow { 
        background-color: #fff3cd; color: #856404; 
        padding: 2px 4px; border-radius: 4px; border: 1px solid #ffeeba; 
        font-weight: bold;
    }

    .highlight-blue { 
        background-color: #d1ecf1; color: #0c5460; 
        padding: 2px 4px; border-radius: 4px; border: 1px solid #bee5eb; font-weight: bold; 
    }

    .topico-item {
        display: block;
        margin-left: 20px;
        margin-bottom: 4px;
        text-indent: -15px; 
    }

    .border-ok { border-left: 6px solid #28a745 !important; }
    .border-warn { border-left: 6px solid #ffc107 !important; } 
    .border-info { border-left: 6px solid #17a2b8 !important; }

    div[data-testid="stMetric"] {
        background-color: #f8f9fa; border: 1px solid #dee2e6; padding: 10px; border-radius: 5px; text-align: center;
    }

    section[data-testid="stSidebar"] {
        display: block !important;
        visibility: visible !important;
        width: 250px !important;
        min-width: 250px !important;
        max-width: 250px !important;
        margin-left: 0 !important;
        transform: translateX(0) !important;
        transition: none !important;
        position: relative !important;
        background-color: #f0f2f6 !important;
        z-index: 999 !important;
    }

    section[data-testid="stSidebar"] > div:first-child {
        width: 250px !important;
        min-width: 250px !important;
    }

    section[data-testid="stSidebar"][aria-expanded="false"],
    section[data-testid="stSidebar"][aria-expanded="true"] {
        margin-left: 0 !important;
        transform: translateX(0) !important;
    }

    button[kind="header"],
    [data-testid="collapsedControl"],
    button[data-testid="baseButton-header"] {
        display: none !important;
    }
</style>
""", unsafe_allow_html=True)

# ----------------- 2. CONFIGURAÇÃO -----------------
# Mantive a lista padrão curta (não é mais a fonte primária), o novo código tenta descobrir modelos via ListModels.
MODELOS_PARA_TENTAR = [
    "gemini-1.5-flash",
    "gemini-1.5",
    "text-bison-001",
    "chat-bison-001"
]

SECOES_PACIENTE = [
    "CABEÇALHO DA BULA",
    "APRESENTAÇÕES",
    "COMPOSIÇÃO",
    "PARA QUE ESTE MEDICAMENTO É INDICADO?",
    "COMO ESTE MEDICAMENTO FUNCIONA?",
    "QUANDO NÃO DEVO USAR ESTE MEDICAMENTO?",
    "O QUE DEVO SABER ANTES DE USAR ESTE MEDICAMENTO?",
    "ONDE, COMO E POR QUANTO TEMPO POSSO GUARDAR ESTE MEDICAMENTO?",
    "COMO DEVO USAR ESTE MEDICAMENTO?",
    "O QUE DEVO FAZER QUANDO EU ME ESQUECER DE USAR ESTE MEDICAMENTO?",
    "QUAIS OS MALES QUE ESTE MEDICAMENTO PODE CAUSAR?",
    "O QUE FAZER SE ALGUÉM USAR UMA QUANTIDADE MAIOR DO QUE A INDICADA DESTE MEDICAMENTO?",
    "DIZERES LEGAIS"
]

SECOES_PROFISSIONAL = []

SIMILARITY_THRESHOLD = 0.92  # mais rígido para reduzir falsos-positivos

# ----------------- helpers para descoberta e uso de modelos -----------------

def safe_list_models():
    """
    Tenta chamar genai.list_models() e retorna lista simples de nomes (strings).
    Trata formatos diversos de retorno.
    """
    try:
        res = genai.list_models()
    except Exception as e:
        # se list_models não existir ou falhar, propaga a exceção para o chamador
        raise

    model_names = []
    # vários formatos possíveis: objeto com 'models' list, ou lista direta
    if isinstance(res, dict):
        candidates = []
        if "models" in res and isinstance(res["models"], list):
            candidates = res["models"]
        elif "data" in res and isinstance(res["data"], list):
            candidates = res["data"]
        else:
            # tentar interpretar por tentativa
            candidates = []
        for m in candidates:
            if isinstance(m, dict):
                # formatos possíveis: {'name': '...'} ou {'model': '...'}
                name = m.get("name") or m.get("model") or m.get("id")
                if name:
                    model_names.append(name)
            elif isinstance(m, str):
                model_names.append(m)
    elif isinstance(res, list):
        for m in res:
            if isinstance(m, str):
                model_names.append(m)
            elif isinstance(m, dict):
                name = m.get("name") or m.get("model") or m.get("id")
                if name:
                    model_names.append(name)
    else:
        # objeto custom - tentar acessar atributo 'models'
        try:
            candidates = getattr(res, "models", None) or getattr(res, "data", None)
            if candidates:
                for m in candidates:
                    if isinstance(m, dict):
                        name = m.get("name") or m.get("model") or m.get("id")
                        if name:
                            model_names.append(name)
        except Exception:
            pass

    # única filtragem: remover None e duplicatas, manter ordem
    seen = set()
    final = []
    for n in model_names:
        if not n:
            continue
        if n in seen:
            continue
        seen.add(n)
        final.append(n)
    return final

def extract_text_from_genai_response(response_obj):
    """
    Extrai texto de várias formas de resposta possíveis:
     - response.text
     - response.candidates[0].content
     - response['candidates'][0]['content']
     - response['output'][0]['content']
     - response.get('output')
    Retorna string ou None.
    """
    # 1) atributo .text
    try:
        txt = getattr(response_obj, "text", None)
        if txt:
            return txt
    except Exception:
        pass

    # 2) atributo .candidates (obj or dict)
    try:
        candidates = getattr(response_obj, "candidates", None)
        if candidates and len(candidates) > 0:
            first = candidates[0]
            if isinstance(first, dict):
                return first.get("content") or first.get("output") or first.get("text")
            else:
                # objeto com atributo 'content'?
                return getattr(first, "content", None) or getattr(first, "text", None)
    except Exception:
        pass

    # 3) dict-like
    try:
        if isinstance(response_obj, dict):
            if "candidates" in response_obj and response_obj["candidates"]:
                c0 = response_obj["candidates"][0]
                if isinstance(c0, dict):
                    return c0.get("content") or c0.get("output") or c0.get("text")
            if "output" in response_obj and response_obj["output"]:
                out0 = response_obj["output"][0]
                if isinstance(out0, dict):
                    return out0.get("content") or out0.get("text")
            # alguns retornos têm 'result' -> 'content'
            if "result" in response_obj:
                r = response_obj["result"]
                if isinstance(r, dict):
                    return r.get("content") or r.get("text")
    except Exception:
        pass

    # 4) tentar str()
    try:
        s = str(response_obj)
        if s and len(s) > 0:
            return s
    except Exception:
        pass

    return None

def choose_candidate_models_from_list(model_list):
    """
    Recebe uma lista de nomes de modelos e ordena/filtra candidatos prováveis para geração.
    Prioriza nomes que contenham gemini, bison, text, chat.
    """
    priorities = []
    for m in model_list:
        lname = m.lower()
        score = 0
        if "gemini" in lname:
            score += 100
        if "bison" in lname:
            score += 80
        if "text" in lname:
            score += 60
        if "chat" in lname:
            score += 50
        # preferir modelos sem '-preview' ou '-beta' no nome
        if "-preview" in lname or "-beta" in lname:
            score -= 30
        priorities.append((score, m))
    priorities.sort(reverse=True, key=lambda x: x[0])
    return [m for _, m in priorities]

# ----------------- 3..9: o resto do código original (sem alterações lógicas substanciais) -----------------
# Para brevidade aqui eu mantenho o mesmo código do seu app: limpeza, extração, comparação, UI, etc.
# (copiei integralmente suas funções de limpeza/extracao/diff/visual para manter comportamento idêntico)

# ----------------- 3. LIMPEZA AUTOMATIZADA (METADADOS E RODAPÉS) -----------------

def clean_metadata_and_footers(texto: str) -> str:
    if not texto:
        return texto
    t = texto

    # nomes de arquivo longos em MAIÚSCULAS/underscores
    t = re.sub(r'(?m)^\s*[A-Z0-9_]{8,}\s*$', '', t)

    # linhas com metadados específicos
    patterns_line = [
        r'(?im)^\s*.*medida\s+da\s+bula.*$',
        r'(?im)^\s*.*tipologia\s+da\s+bula.*$',
        r'(?im)^\s*.*tipologia:.*$',
        r'(?im)^\s*.*impress(ã|a)o.*:.*$',
        r'(?im)^\s*.*papel:.*$',
        r'(?im)^\s*.*cor:.*$',
        r'(?im)^\s*.*frente\/verso.*$',
    ]
    for p in patterns_line:
        t = re.sub(p, '', t)

    # dimensões e paginação
    t = re.sub(r'(?im)\d{1,2},\d{2}\s*cm\s*[x×X]\s*\d{1,2},\d{2}\s*cm', '', t)
    page_patterns = [
        r'(?im)\bBula(?:\s+ao\s+Paciente)?\s+P[aá]gina\s*\d+\s*(?:de|\/)\s*\d+\b',
        r'(?im)\bBula(?:\s+ao\s+Paciente)?\s+P[aá]gina\s*\d+\b',
        r'(?im)\bP[aá]gina\s*\d+\s*(?:de|\/)\s*\d+\b',
        r'(?im)\bP[aá]gina\s*\d+\b'
    ]
    for p in page_patterns:
        t = re.sub(p, '', t)

    t = re.sub(r'(?im)\bfrente\b', '', t)
    t = re.sub(r'(?im)\bverso\b', '', t)
    t = re.sub(r'(?im)\bBUL[_A-Z0-9-]*\b', '', t)

    # remover quebras por hífen (quebra de sílaba no final de linha)
    t = re.sub(r'-\s*\n\s*', '', t)

    # normalizações simples de espaçamento
    t = re.sub(r'[ \t]{2,}', ' ', t)
    t = re.sub(r'\r', '\n', t)
    t = re.sub(r'\n{3,}', '\n\n', t)

    # remover linhas vazias extras
    lines = [ln.rstrip() for ln in t.splitlines()]
    lines = [ln for ln in lines if ln.strip() != ""]
    t = "\n".join(lines)

    return t.strip()

# ----------------- 4. LOCALIZAÇÃO DE TÍTULOS E EXTRAÇÃO DE CABEÇALHO (SEGURA) -----------------

def build_section_pattern(title: str) -> str:
    words = re.findall(r'\w+', title, flags=re.UNICODE)
    if not words:
        return None
    # pattern tolerante: permite qualquer não-alfanumérico entre as palavras
    pattern = r'\b' + r'\W+'.join(map(re.escape, words)) + r'\b'
    return pattern

def find_first_section_index(texto: str, section_titles: list) -> int:
    menor = None
    for title in section_titles:
        if not title:
            continue
        patt = build_section_pattern(title)
        if not patt:
            continue
        m = re.search(patt, texto, flags=re.IGNORECASE | re.UNICODE)
        if m:
            idx = m.start()
            if menor is None or idx < menor:
                menor = idx
    return -1 if menor is None else menor

def safe_extract_header(texto: str, secoes_alvo: list) -> str:
    if not texto:
        return ""
    idx = find_first_section_index(texto, [s for s in secoes_alvo if s.strip().upper() != "CABEÇALHO DA BULA"])
    if idx == -1:
        # fallback: procurar 'APRESENTA' isoladamente
        m = re.search(r'\bAPRESENTA\S*\b', texto, flags=re.IGNORECASE)
        idx = m.start() if m else -1
    if idx == -1:
        return ""
    header_raw = texto[:idx].strip()
    header_raw = re.sub(r'(?m)^\s*[IVXLCDM]+\s*[–-]\s*', '', header_raw)
    header_clean = clean_metadata_and_footers(header_raw)
    if len(header_clean) < 20:
        return ""
    if len(header_clean) > max(2000, int(len(texto) * 0.35)):
        return ""
    return header_clean.strip()

# ----------------- 4b. EXTRAIR SEÇÃO DO TEXTO BRUTO (FALLBACK MELHORADO) -----------------

def _build_flexible_title_regex(title: str):
    """Cria regex que aceita numeração opcional antes do título e pontuações."""
    words = re.findall(r'\w+', title, flags=re.UNICODE)
    if not words:
        return None
    core = r'\W+'.join(map(re.escape, words))
    # aceita opcionalmente "2. ", "2) ", "II - " antes do título, ou começo de linha
    regex = rf'(?:^|\n)\s*(?:\d{{1,2}}\s*[\.\)\-]\s*|[IVXLCDM]+\s*[–-]\s*)?{core}'
    return regex

def extract_section_from_raw(texto: str, section_title: str, sections_list: list) -> str:
    if not texto or not section_title:
        return ""

    patt = _build_flexible_title_regex(section_title)
    if patt:
        m = re.search(patt, texto, flags=re.IGNORECASE | re.UNICODE)
    else:
        m = re.search(build_section_pattern(section_title), texto, flags=re.IGNORECASE | re.UNICODE)

    if not m:
        keywords = re.findall(r'\w+', section_title)
        if keywords:
            for kcount in (min(3, len(keywords)), len(keywords)):
                core = r'\W+'.join(map(re.escape, keywords[:kcount]))
                m = re.search(core, texto, flags=re.IGNORECASE | re.UNICODE)
                if m:
                    break
    if not m:
        return ""

    start = m.end()

    menor = None
    for s in sections_list:
        if not s:
            continue
        if s.strip().upper() == section_title.strip().upper():
            continue
        p2 = _build_flexible_title_regex(s)
        if not p2:
            p2 = build_section_pattern(s)
        m2 = re.search(p2, texto[start:], flags=re.IGNORECASE | re.UNICODE)
        if m2:
            idx = start + m2.start()
            if menor is None or idx < menor:
                menor = idx

    if menor is None:
        m3 = re.search(r'\n{2,}([A-ZÀ-Ý0-9 \-]{6,})\n', texto[start:])
        if m3:
            candidate = m3.group(1).strip()
            if len(candidate.split()) <= 6:
                menor = start + m3.start()

    end = menor if menor is not None else len(texto)

    section_raw = texto[start:end].strip()

    section_raw = re.sub(r'(?m)^\s*\d+\s*[\.\)\-]?\s*', '', section_raw)
    section_raw = re.sub(r'(?m)^\s*[IVXLCDM]+\s*[–-]\s*', '', section_raw)

    section_clean = re.sub(r'\r', '\n', section_raw).strip()

    if len(re.sub(r'<[^>]+>', '', section_clean).strip()) < 10:
        return ""
    if len(re.sub(r'<[^>]+>', '', section_clean)) > max(8000, int(len(texto) * 0.8)):
        return ""

    return section_clean

# ----------------- 5. NORMALIZAÇÃO AVANÇADA PARA COMPARAÇÃO -----------------

def strip_accents(s: str) -> str:
    return unicodedata.normalize('NFKD', s).encode('ASCII', 'ignore').decode('ASCII')

def tokenize_words(s: str):
    return re.findall(r'\w+', s, flags=re.UNICODE)

def remove_section_titles_for_comparison(texto: str, sections_list=SECOES_PACIENTE) -> str:
    if not texto:
        return texto
    t = texto
    for s in sections_list:
        if not s or s.strip().upper() == "CABEÇALHO DA BULA":
            continue
        patt = build_section_pattern(s)
        if patt:
            t = re.sub(patt, ' ', t, flags=re.IGNORECASE)
    t = re.sub(r'(?im)\bInform[aç]o(?:es)?\s+ao\s+paciente\b', ' ', t)
    t = re.sub(r'\s{2,}', ' ', t)
    return t.strip()

def normalize_for_comparison(text: str) -> str:
    if not text:
        return ""
    t = clean_metadata_and_footers(text)
    t = remove_section_titles_for_comparison(t)
    t = re.sub(r'(?m)^\s*\d+\s*[\.\)\-]\s*', '', t)
    t = re.sub(r'\b\d+\s*[\.\)]\s+', ' ', t)
    t = re.sub(r'(?m)^\s*[IVXLCDM]+\s*[–-]\s*', '', t)
    t = re.sub(r'<[^>]+>', '', t)
    t = strip_accents(t).lower()
    t = re.sub(r'[^a-z0-9\s]', ' ', t)
    t = re.sub(r'\s+', ' ', t).strip()
    return t

def jaccard_similarity(a: str, b: str) -> float:
    sa = set(tokenize_words(a))
    sb = set(tokenize_words(b))
    if not sa and not sb:
        return 1.0
    if not sa or not sb:
        return 0.0
    inter = sa.intersection(sb)
    union = sa.union(sb)
    return len(inter) / len(union)

# ----------------- 6. ORTOGRAFIA E VISUAL -----------------

def verificar_ortografia_inteligente(texto):
    try:
        spell = SpellChecker(language='pt')
        whitelist = {'mg','ml','mcg','ui','g','kg','l','dl','mmhg','bpm','kcal','anvisa','cnpj','cep','sac','bula'}
        spell.word_frequency.load_words(whitelist)
        tokens = re.split(r'(<[^>]+>|\s+|[().,:;!?/\[\]])', texto)
        resultado = []
        for token in tokens:
            if not token.strip() or token.startswith('<') or not any(c.isalpha() for c in token):
                resultado.append(token); continue
            palavra_limpa = re.sub(r'[^a-zA-ZáàâãéèêíïóôõöúçñÁÀÂÃÉÈÊÍÏÓÔÕÖÚÇÑ-]', '', token)
            if (not palavra_limpa or len(palavra_limpa) < 4 or any(c.isdigit() for c in token) or '-' in palavra_limpa or palavra_limpa[0].isupper()):
                resultado.append(token); continue
            resultado.append(token)
        return "".join(resultado)
    except:
        return texto

def melhorar_visual_topicos(texto_html):
    linhas = re.split(r'(<br>|\n)', texto_html)
    novo_texto = []
    for linha in linhas:
        if re.search(r'^\s*[-•*]\s+', re.sub(r'<[^>]+>', '', linha).strip()):
            linha_limpa = re.sub(r'^\s*[-•*]\s+', '', linha)
            novo_texto.append(f'<div class="topico-item">• {linha_limpa}</div>')
        else:
            novo_texto.append(linha)
    return "".join(novo_texto)

def destacar_datas(texto):
    padrao = r'(Esta\s+bula\s+foi\s+(?:atualizada\s+conforme\s+Bula\s+Padrão\s+)?aprovada\s+pela\s+Anvisa\s+em\s*)(\d{2}/\d{2}/\d{4}|\d{2}/\d{4})'
    def replacer(match):
        return f'{match.group(1)}<span class="highlight-blue">{match.group(2)}</span>'
    return re.sub(padrao, replacer, texto, count=0, flags=re.IGNORECASE | re.DOTALL)

# ----------------- 7. DIFF E REGRAS DE DECISÃO -----------------

def diff_palavra_a_palavra(texto_ref, texto_novo):
    ref_sem_tags = re.sub(r'<[^>]+>', '', texto_ref)
    novo_sem_tags = re.sub(r'<[^>]+>', '', texto_novo)
    palavras_ref = ref_sem_tags.split()
    palavras_novo = novo_sem_tags.split()
    matcher = difflib.SequenceMatcher(None, palavras_ref, palavras_novo)
    html_ref_list = []
    html_novo_list = []
    tem_diff = False
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            texto = " ".join(palavras_ref[i1:i2]); html_ref_list.append(texto); html_novo_list.append(texto)
        elif tag == 'replace':
            html_ref_list.append(f'<span class="highlight-yellow">{" ".join(palavras_ref[i1:i2])}</span>')
            html_novo_list.append(f'<span class="highlight-yellow">{" ".join(palavras_novo[j1:j2])}</span>')
            tem_diff = True
        elif tag == 'delete':
            html_ref_list.append(f'<span class="highlight-yellow">{" ".join(palavras_ref[i1:i2])}</span>')
            tem_diff = True
        elif tag == 'insert':
            html_novo_list.append(f'<span class="highlight-yellow">{" ".join(palavras_novo[j1:j2])}</span>')
            tem_diff = True
    return " ".join(html_ref_list), " ".join(html_novo_list), tem_diff

def gerar_diff_html(texto_ref, texto_novo, secoes_alvo=SECOES_PACIENTE):
    if not texto_ref: texto_ref = ""
    if not texto_novo: texto_novo = ""

    comp_ref = clean_metadata_and_footers(texto_ref)
    comp_novo = clean_metadata_and_footers(texto_novo)

    norm_ref = normalize_for_comparison(comp_ref)
    norm_novo = normalize_for_comparison(comp_novo)

    if norm_ref == norm_novo:
        html_ref = comp_ref.replace('\n', '<br>'); html_ref = melhorar_visual_topicos(html_ref)
        html_novo = verificar_ortografia_inteligente(comp_novo); html_novo = html_novo.replace('\n', '<br>'); html_novo = melhorar_visual_topicos(html_novo)
        return html_ref, html_novo, False

    if norm_ref and norm_novo:
        shorter, longer = (norm_ref, norm_novo) if len(norm_ref) <= len(norm_novo) else (norm_novo, norm_ref)
        if shorter and shorter in longer:
            prop = len(shorter) / max(1, len(longer))
            if prop >= 0.88:
                html_ref = comp_ref.replace('\n', '<br>'); html_ref = melhorar_visual_topicos(html_ref)
                html_novo = verificar_ortografia_inteligente(comp_novo); html_novo = html_novo.replace('\n', '<br>'); html_novo = melhorar_visual_topicos(html_novo)
                return html_ref, html_novo, False

    ratio = difflib.SequenceMatcher(None, norm_ref, norm_novo).ratio()
    jacc = jaccard_similarity(norm_ref, norm_novo)
    if ratio >= SIMILARITY_THRESHOLD or jacc >= SIMILARITY_THRESHOLD:
        html_ref = comp_ref.replace('\n', '<br>'); html_ref = melhorar_visual_topicos(html_ref)
        html_novo = verificar_ortografia_inteligente(comp_novo); html_novo = html_novo.replace('\n', '<br>'); html_novo = melhorar_visual_topicos(html_novo)
        return html_ref, html_novo, False

    r_html, n_html, diff_bool = diff_palavra_a_palavra(comp_ref, comp_novo)
    n_html_final = verificar_ortografia_inteligente(n_html); n_html_final = melhorar_visual_topicos(n_html_final)
    r_html_final = melhorar_visual_topicos(r_html.replace('\n', '<br>'))
    return r_html_final, n_html_final, diff_bool

# ----------------- 8. EXTRAÇÃO DE TEXTO -----------------

def extract_text_from_file(uploaded_file):
    try:
        text = ""
        name = getattr(uploaded_file, "name", "uploaded")
        if name.lower().endswith('.pdf'):
            uploaded_file.seek(0)
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            for page in doc:
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                for block in blocks:
                    if block.get("type") != 0:
                        continue
                    block_text = ""
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            content = span.get("text", "")
                            if not content.strip():
                                line_text += content; continue
                            flags = span.get("flags", 0)
                            font_name = span.get("font", "").lower()
                            is_bold = ((flags & 16) or "bold" in font_name or "black" in font_name or "heavy" in font_name or "semibold" in font_name)
                            is_italic = ((flags & 2) or "italic" in font_name or "oblique" in font_name)
                            formatted_text = content
                            if is_bold and is_italic:
                                formatted_text = f"<b><i>{content}</i></b>"
                            elif is_bold:
                                formatted_text = f"<b>{content}</b>"
                            elif is_italic:
                                formatted_text = f"<i>{content}</i>"
                            line_text += formatted_text
                        block_text += line_text.strip() + " "
                    text += block_text.strip() + "\n\n"
            doc.close()
        elif name.lower().endswith('.docx'):
            try:
                uploaded_file.seek(0)
                doc = docx.Document(uploaded_file)
            except Exception:
                uploaded_file.seek(0)
                doc = docx.Document(BytesIO(uploaded_file.read()))
            for para in doc.paragraphs:
                para_text = ""
                for run in para.runs:
                    content = run.text
                    if not content: continue
                    is_bold = run.bold is True
                    is_italic = run.italic is True
                    formatted_text = content
                    if is_bold and is_italic:
                        formatted_text = f"<b><i>{content}</i></b>"
                    elif is_bold:
                        formatted_text = f"<b>{content}</b>"
                    elif is_italic:
                        formatted_text = f"<i>{content}</i>"
                    para_text += formatted_text
                text += para_text + "\n\n"
        return clean_metadata_and_footers(text.strip())
    except Exception as e:
        st.error(f"Erro ao extrair texto do arquivo {getattr(uploaded_file, 'name', '')}: {str(e)}")
        return ""

# ----------------- 9. UI PRINCIPAL E FLUXO -----------------
st.title("💊 Conferência MKT")

tipo_bula = st.radio("Escolha o Tipo de Bula:", ("Paciente",), horizontal=True)

c1, c2 = st.columns(2)
f1 = c1.file_uploader("📜 Bula BELFAR", type=["pdf", "docx"], key="f1")
f2 = c2.file_uploader("📜 Bula MKT", type=["pdf", "docx"], key="f2")

if st.button("🚀 Processar Conferência", key="process_button"):
    st.info("Iniciando processamento...")
    if not f1 or not f2:
        st.warning("Por favor, envie ambos os arquivos antes de processar.")
        st.stop()
    else:
        st.info(f"Arquivo BELFAR detectado: {getattr(f1, 'name', 'desconhecido')}")
        st.info(f"Arquivo MKT detectado: {getattr(f2, 'name', 'desconhecido')}")

    keys_raw = [
        st.secrets.get("GEMINI_API_KEY"),
        st.secrets.get("GEMINI_API_KEY2"),
        st.secrets.get("GEMINI_API_KEY3")
    ]
    keys_validas = [k for k in keys_raw if k]

    if not keys_validas:
        st.error("Erro crítico: nenhuma API key encontrada. Sem fallback definido - impossível prosseguir.")
        st.stop()

    if f1 and f2:
        secoes_alvo = SECOES_PACIENTE if tipo_bula == "Paciente" else SECOES_PROFISSIONAL

        with st.spinner("Lendo arquivos e conectando à IA..."):
            try:
                f1.seek(0); f2.seek(0)
                t_anvisa = extract_text_from_file(f1)  # já limpo de metadados
                t_mkt = extract_text_from_file(f2)
                st.info(f"Tamanho do texto BELFAR (após limpeza): {len(t_anvisa)}")
                st.info(f"Tamanho do texto MKT (após limpeza): {len(t_mkt)}")
            except Exception as e:
                st.exception(f"Falha ao extrair textos: {e}")
                st.stop()

            if len(t_anvisa) < 20 or len(t_mkt) < 20:
                st.error("Arquivo vazio ou ilegível."); st.stop()

            prompt = f"""
            Você é um Extrator de Dados Farmacêuticos Rigoroso.

            INPUT TEXTO 1 (REF): {t_anvisa[:150000]}
            INPUT TEXTO 2 (MKT): {t_mkt[:150000]}

            REGRAS CRÍTICAS DE EXTRAÇÃO:

**CABEÇALHO DA BULA**: Extrair TODO o conteúdo desde o início do documento ATÉ (mas NÃO incluindo) o título "APRESENTAÇÕES". 
IMPORTANTE: Remover APENAS os algarismos romanos (I, II, III) e o hífen/traço, mantendo o texto.
PARAR antes de encontrar "APRESENTAÇÕES" ou qualquer variação.

**SEÇÕES NORMAIS**: A partir de "APRESENTAÇÕES" até "DIZERES LEGAIS", extrair cada seção completa.

**FORMATAÇÃO**: 
Manter <b> e <i> EXATAMENTE como está
NÃO corrigir português
NÃO resumir
Ignorar linhas horizontais/elementos gráficos

Se uma seção não existir, NÃO inclua no JSON.

LISTA DE SEÇÕES ESPERADAS: {secoes_alvo}

SAÍDA JSON:
{{
 "data_anvisa_ref": "dd/mm/aaaa",
 "data_anvisa_mkt": "dd/mm/aaaa",
 "secoes": [
     {{
         "titulo": "NOME EXATO DA SEÇÃO",
         "texto_anvisa": "conteúdo completo",
         "texto_mkt": "conteúdo completo"
     }}
 ]
}}
"""

            response = None
            sucesso = False
            log_erros = []
            modelos_testados = []

            # tentativa: para cada key, descubro modelos (list_models) e tento os mais prováveis
            for idx_key, key in enumerate(keys_validas):
                if sucesso:
                    break
                try:
                    genai.configure(api_key=key)
                except Exception as e:
                    log_erros.append(f"Key {idx_key+1} | configure: {str(e)}")
                    continue

                # 1) tentar modelos estáticos rápidos (se quiser manter)
                for modelo in MODELOS_PARA_TENTAR:
                    modelos_testados.append((idx_key+1, modelo))
                    try:
                        st.info(f"Tentando modelo estático: {modelo} com key {idx_key+1}")
                        model = genai.GenerativeModel(
                            modelo,
                            generation_config={"response_mime_type": "application/json", "temperature": 0.0}
                        )
                        response = model.generate_content(prompt)
                        sucesso = True
                        st.info(f"Modelo aceito: {modelo}")
                        break
                    except Exception as e:
                        log_erros.append(f"Key {idx_key+1} | {modelo}: {str(e)}")
                        time.sleep(0.2)
                        continue

                if sucesso:
                    break

                # 2) descobrir modelos disponíveis via ListModels / list_models
                try:
                    st.info(f"Chamando ListModels para key {idx_key+1}...")
                    available = safe_list_models()
                    if not available:
                        log_erros.append(f"Key {idx_key+1} | list_models retornou vazio ou formato inesperado.")
                        continue

                    st.info(f"Modelos retornados pela API (key {idx_key+1}): {available[:30]}{'...' if len(available)>30 else ''}")
                    # escolher candidatos ordenados
                    candidates = choose_candidate_models_from_list(available)
                    # garantir que tentamos alguns nomes prioritários
                    tried_local = set()
                    for modelo in candidates:
                        if modelo in tried_local:
                            continue
                        tried_local.add(modelo)
                        modelos_testados.append((idx_key+1, modelo))
                        try:
                            st.info(f"Tentando modelo listado: {modelo} com key {idx_key+1}")
                            # duas formas de chamar: genai.generate_text (se existir) ou GenerativeModel
                            if hasattr(genai, "generate_text"):
                                # tentativa com generate_text (API cliente pode expor isso)
                                response = genai.generate_text(model=modelo, prompt=prompt)
                            else:
                                model = genai.GenerativeModel(
                                    modelo,
                                    generation_config={"response_mime_type": "application/json", "temperature": 0.0}
                                )
                                response = model.generate_content(prompt)
                            sucesso = True
                            st.info(f"Modelo aceito: {modelo}")
                            break
                        except Exception as e:
                            log_erros.append(f"Key {idx_key+1} | {modelo}: {str(e)}")
                            time.sleep(0.2)
                            continue

                    if sucesso:
                        break
                    else:
                        # mostrar os modelos disponíveis para o usuário (para colar aqui)
                        st.warning(f"Key {idx_key+1}: não foi possível usar nenhum modelo automaticamente. Modelos disponíveis listados abaixo.")
                        st.code("\n".join(available))
                except Exception as e:
                    log_erros.append(f"Key {idx_key+1} | list_models/descoberta: {str(e)}")
                    continue

            if not sucesso:
                st.error("❌ Falha total ao chamar a API Gemini/Generative. Sem fallback configurado — abortando.")
                if log_erros:
                    st.code("\n".join(log_erros))
                st.stop()

            # se chegou aqui, temos response da IA
            try:
                if response is None:
                    st.error("Resposta da IA vazia (response is None).")
                    st.stop()
                resp_text = extract_text_from_genai_response(response)
                if not resp_text:
                    st.error("Resposta da IA não contém texto reconhecível.")
                    try:
                        st.code(str(response))
                    except:
                        pass
                    st.stop()
                resultado = json.loads(resp_text)
            except Exception as e:
                st.exception(f"Erro ao decodificar JSON da resposta da IA: {e}")
                try:
                    st.code(resp_text)
                except:
                    pass
                st.stop()

            try:
                data_ref = resultado.get("data_anvisa_ref", "-")
                data_mkt = resultado.get("data_anvisa_mkt", "-")
                dados_secoes = resultado.get("secoes", [])

                secoes_finais = []
                divs_count = 0

                for item in dados_secoes:
                    titulo = item.get('titulo', '').strip()
                    txt_ref = item.get('texto_anvisa', '').strip()
                    txt_mkt = item.get('texto_mkt', '').strip()

                    if not txt_ref:
                        tentativa = extract_section_from_raw(t_anvisa, titulo, secoes_alvo)
                        if tentativa:
                            txt_ref = tentativa
                    if not txt_mkt:
                        tentativa2 = extract_section_from_raw(t_mkt, titulo, secoes_alvo)
                        if tentativa2:
                            txt_mkt = tentativa2

                    txt_ref = clean_metadata_and_footers(txt_ref)
                    txt_mkt = clean_metadata_and_footers(txt_mkt)

                    if "CABEÇALHO" in titulo.upper():
                        novo_ref = safe_extract_header(t_anvisa, secoes_alvo)
                        if novo_ref and (not txt_ref or len(novo_ref) < len(txt_ref) or len(txt_ref) < 50):
                            txt_ref = novo_ref

                        novo_mkt = safe_extract_header(t_mkt, secoes_alvo)
                        if novo_mkt and (not txt_mkt or len(novo_mkt) < len(txt_mkt) or len(txt_mkt) < 50):
                            txt_mkt = novo_mkt

                        txt_ref = re.sub(r'(?m)^\s*[IVXLCDM]+\s*[–-]\s*', '', txt_ref)
                        txt_mkt = re.sub(r'(?m)^\s*[IVXLCDM]+\s*[–-]\s*', '', txt_mkt)

                    if "DIZERES LEGAIS" in titulo.upper():
                        html_ref = destacar_datas(txt_ref).replace('\n', '<br>')
                        html_novo = destacar_datas(txt_mkt).replace('\n', '<br>')
                        html_ref = verificar_ortografia_inteligente(html_ref)
                        html_novo = verificar_ortografia_inteligente(html_novo)
                        html_ref = melhorar_visual_topicos(html_ref)
                        html_novo = melhorar_visual_topicos(html_novo)
                        status = "CONFORME"
                    else:
                        html_ref, html_novo, teve_diff = gerar_diff_html(txt_ref, txt_mkt, secoes_alvo)
                        status = "DIVERGENTE" if teve_diff else "CONFORME"
                        if teve_diff:
                            divs_count += 1

                    secoes_finais.append({
                        "titulo": titulo,
                        "texto_anvisa": html_ref,
                        "texto_mkt": html_novo,
                        "status": status
                    })

                st.markdown("### 📊 Resumo")
                c1, c2, c3 = st.columns(3)
                c1.metric("Data BELFAR", data_ref)
                c2.metric("Data MKT", data_mkt, delta="Igual" if data_ref == data_mkt else "Diferente")
                c3.metric("Seções", len(secoes_finais))

                sub1, sub2 = st.columns(2)
                sub1.info(f"✅ Conformes: {len(secoes_finais) - divs_count}")
                if divs_count > 0:
                    sub2.warning(f"⚠️ Divergentes: {divs_count}")
                else:
                    sub2.success("✨ Divergências: 0")

                st.divider()

                for item in secoes_finais:
                    status = item['status']
                    titulo = item['titulo']

                    if "DIZERES LEGAIS" in titulo.upper():
                        icon, css, aberto = "⚖️", "border-info", False
                    elif status == "CONFORME":
                        icon, css, aberto = "✅", "border-ok", False
                    else:
                        icon, css, aberto = "⚠️", "border-warn", True

                    with st.expander(f"{icon} {titulo}", expanded=aberto):
                        ce, cd = st.columns(2)
                        with ce:
                            st.caption("BELFAR")
                            st.markdown(f'<div class="texto-box {css}">{item["texto_anvisa"]}</div>', unsafe_allow_html=True)
                        with cd:
                            st.caption("MKT")
                            st.markdown(f'<div class="texto-box {css}">{item["texto_mkt"]}</div>', unsafe_allow_html=True)

            except Exception as e:
                st.exception(f"Erro ao processar resultado: {e}")
                try:
                    st.code(str(response))
                except:
                    pass
                st.stop()

else:
    st.info("Aguardando ação. Adicione os arquivos e clique em 'Processar Conferência'.")
