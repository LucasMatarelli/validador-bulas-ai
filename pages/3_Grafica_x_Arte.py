# app.py - Gráfica x Arte (Versão Final com Diff Ajustado)
import streamlit as st
import google.generativeai as genai
import fitz  # PyMuPDF
import docx
import json
import difflib
import re
import unicodedata
import time
from spellchecker import SpellChecker
from io import BytesIO

# ----------------- 1. VISUAL & CSS -----------------
st.set_page_config(page_title="Gráfica x Arte", page_icon="💊", layout="wide")

st.markdown("""
<style>
    [data-testid="stHeader"] { visibility: hidden; }

    .texto-box { 
        font-family: 'Segoe UI', sans-serif;
        font-size: 0.95rem;
        line-height: 1.7;
        color: #212529;
        background-color: #ffffff;
        padding: 25px;
        border-radius: 8px;
        border: 1px solid #ced4da;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        text-align: left;
        height: 100%;
        overflow-y: auto;
        max-height: 600px;
    }

    .highlight-yellow { 
        background-color: #fff3cd; color: #856404; 
        padding: 2px 4px; border-radius: 4px; border: 1px solid #ffeeba; 
        font-weight: bold;
    }

    .highlight-blue { 
        background-color: #d1ecf1; color: #0c5460; 
        padding: 2px 4px; border-radius: 4px; border: 1px solid #bee5eb; font-weight: bold; 
    }

    .topico-item {
        display: block;
        margin-left: 20px;
        margin-bottom: 4px;
        text-indent: -15px; 
    }

    .border-ok { border-left: 6px solid #28a745 !important; }
    .border-warn { border-left: 6px solid #ffc107 !important; } 
    .border-info { border-left: 6px solid #17a2b8 !important; }

    div[data-testid="stMetric"] {
        background-color: #f8f9fa; border: 1px solid #dee2e6; padding: 10px; border-radius: 5px; text-align: center;
    }
</style>
""", unsafe_allow_html=True)

# ----------------- 2. CONFIGURAÇÃO -----------------
MODELOS_PARA_TENTAR = [
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
    "gemini-1.5-flash"
]

SECOES_PACIENTE = [
    "APRESENTAÇÕES",
    "COMPOSIÇÃO",
    "PARA QUE ESTE MEDICAMENTO É INDICADO?",
    "COMO ESTE MEDICAMENTO FUNCIONA?",
    "QUANDO NÃO DEVO USAR ESTE MEDICAMENTO?",
    "O QUE DEVO SABER ANTES DE USAR ESTE MEDICAMENTO?",
    "ONDE, COMO E POR QUANTO TEMPO POSSO GUARDAR ESTE MEDICAMENTO?",
    "COMO DEVO USAR ESTE MEDICAMENTO?",
    "O QUE DEVO FAZER QUANDO EU ME ESQUECER DE USAR ESTE MEDICAMENTO?",
    "QUAIS OS MALES QUE ESTE MEDICAMENTO PODE CAUSAR?",
    "O QUE FAZER SE ALGUÉM USAR UMA QUANTIDADE MAIOR DO QUE A INDICADA DESTE MEDICAMENTO?",
    "DIZERES LEGAIS"
]

# ----------------- 3. LIMPEZA AUTOMATIZADA -----------------
def clean_metadata_and_footers(texto: str) -> str:
    if not texto: return texto
    t = texto

    t = re.sub(r'(?im)^.*times\s+new\s+roman.*\n?', '', t)
    t = re.sub(r'(?im)^.*negrito.*\n?', '', t)
    t = re.sub(r'(?im)^.*corpo\s*14.*\n?', '', t)
    t = re.sub(r'(?im)^.*\bcontato\b.*\n?', '', t)
    t = re.sub(r'(?i)[\w\.-]+@belfar\.com\.br', '', t)
    t = re.sub(r'(?m)(?:\+?\d{1,3}[-\s]?)?(?:\(?\d{2}\)?[-\s]?)?\d{4,5}[-\s]?\d{4}', '', t)
    t = re.sub(r'(?m)^\s*[A-Z0-9_]{8,}\s*$', '', t)

    patterns_line = [
        r'(?im)^\s*.*medida\s+da\s+bula.*$',
        r'(?im)^\s*.*tipologia\s+da\s+bula.*$',
        r'(?im)^\s*.*tipologia:.*$',
        r'(?im)^\s*.*impress(ã|a)o.*:.*$',
        r'(?im)^\s*.*papel:.*$',
        r'(?im)^\s*.*cor:.*$',
        r'(?im)^\s*.*frente\/verso.*$',
    ]
    for p in patterns_line: t = re.sub(p, '', t)

    t = re.sub(r'(?im)\d{1,2},\d{2}\s*cm\s*[x×X]\s*\d{1,2},\d{2}\s*cm', '', t)
    page_patterns = [
        r'(?im)\bBula(?:\s+ao\s+Paciente)?\s+P[aá]gina\s*\d+\s*(?:de|\/)\s*\d+\b',
        r'(?im)\bBula(?:\s+ao\s+Paciente)?\s+P[aá]gina\s*\d+\b',
        r'(?im)\bP[aá]gina\s*\d+\s*(?:de|\/)\s*\d+\b',
        r'(?im)\bP[aá]gina\s*\d+\b'
    ]
    for p in page_patterns: t = re.sub(p, '', t)

    t = re.sub(r'(?im)\bfrente\b', '', t)
    t = re.sub(r'(?im)\bverso\b', '', t)
    t = re.sub(r'(?im)\bBUL[_A-Z0-9-]*\b', '', t)
    t = re.sub(r'-\s*\n\s*', '', t)
    t = re.sub(r'[ \t]{2,}', ' ', t)
    t = re.sub(r'\r', '\n', t)
    t = re.sub(r'\n{3,}', '\n\n', t)

    lines = [ln.rstrip() for ln in t.splitlines()]
    lines = [ln for ln in lines if ln.strip() != ""]
    t = "\n".join(lines)
    return t.strip()

# ----------------- 4. HELPERS DE TEXTO E REGEX -----------------
def build_section_pattern(title: str) -> str:
    words = re.findall(r'\w+', title, flags=re.UNICODE)
    if not words: return None
    pattern = r'\b' + r'\W+'.join(map(re.escape, words)) + r'\b'
    return pattern

def _build_flexible_title_regex(title: str):
    words = re.findall(r'\w+', title, flags=re.UNICODE)
    if not words: return None
    core = r'\W+'.join(map(re.escape, words))
    regex = rf'(?:^|\n)\s*(?:\d{{1,2}}\s*[\.\)\-]\s*|[IVXLCDM]+\s*[–-]\s*)?{core}'
    return regex

def extract_section_from_raw(texto: str, section_title: str, sections_list: list) -> str:
    if not texto or not section_title: return ""
    patt = _build_flexible_title_regex(section_title)
    if patt:
        m = re.search(patt, texto, flags=re.IGNORECASE | re.UNICODE)
    else:
        m = re.search(build_section_pattern(section_title), texto, flags=re.IGNORECASE | re.UNICODE)
    if not m:
        keywords = re.findall(r'\w+', section_title)
        if keywords:
            for kcount in (min(3, len(keywords)), len(keywords)):
                core = r'\W+'.join(map(re.escape, keywords[:kcount]))
                m = re.search(core, texto, flags=re.IGNORECASE | re.UNICODE)
                if m: break
    if not m: return ""
    start = m.end()
    menor = None
    for s in sections_list:
        if not s: continue
        if s.strip().upper() == section_title.strip().upper(): continue
        p2 = _build_flexible_title_regex(s)
        if not p2: p2 = build_section_pattern(s)
        m2 = re.search(p2, texto[start:], flags=re.IGNORECASE | re.UNICODE)
        if m2:
            idx = start + m2.start()
            if menor is None or idx < menor: menor = idx
    if menor is None:
        m3 = re.search(r'\n{2,}([A-ZÀ-Ý0-9 \-]{6,})\n', texto[start:])
        if m3:
            candidate = m3.group(1).strip()
            if len(candidate.split()) <= 6: menor = start + m3.start()
    end = menor if menor is not None else len(texto)
    section_raw = texto[start:end].strip()
    section_raw = re.sub(r'(?m)^\s*\d+\s*[\.\)\-]?\s*', '', section_raw)
    section_raw = re.sub(r'(?m)^\s*[IVXLCDM]+\s*[–-]\s*', '', section_raw)
    section_clean = re.sub(r'\r', '\n', section_raw).strip()
    if len(re.sub(r'<[^>]+>', '', section_clean).strip()) < 10: return ""
    if len(re.sub(r'<[^>]+>', '', section_clean)) > max(8000, int(len(texto) * 0.8)): return ""
    return section_clean

def convert_markdown_bold_to_html(text: str) -> str:
    if not text: return text
    return re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)

# ----------------- 5. HELPERS DE COMPARAÇÃO (Novos) -----------------
def normalizacao_nuclear(texto: str) -> str:
    """Remove tudo: acentos, tags html, pontuação e espaços, para checar conteúdo puro."""
    if not texto: return ""
    # Remove tags HTML
    t = re.sub(r'<[^>]+>', '', texto)
    # Remove acentos
    t = unicodedata.normalize('NFKD', t).encode('ASCII', 'ignore').decode('ASCII')
    # Mantém apenas letras e números, tudo minúsculo
    t = re.sub(r'[^a-z0-9]', '', t.lower())
    return t

# ----------------- 6. ORTOGRAFIA E VISUAL -----------------
def verificar_ortografia_inteligente(texto):
    try:
        spell = SpellChecker(language='pt')
        whitelist = {'mg','ml','mcg','ui','g','kg','l','dl','mmhg','bpm','kcal','anvisa','cnpj','cep','sac','bula'}
        spell.word_frequency.load_words(whitelist)
        tokens = re.split(r'(<[^>]+>|\s+|[().,:;!?/\[\]])', texto)
        resultado = []
        for token in tokens:
            if not token.strip() or token.startswith('<') or not any(c.isalpha() for c in token):
                resultado.append(token); continue
            palavra_limpa = re.sub(r'[^a-zA-ZáàâãéèêíïóôõöúçñÁÀÂÃÉÈÊÍÏÓÔÕÖÚÇÑ-]', '', token)
            if (not palavra_limpa or len(palavra_limpa) < 4 or any(c.isdigit() for c in token) or '-' in palavra_limpa or palavra_limpa[0].isupper()):
                resultado.append(token); continue
            resultado.append(token)
        return "".join(resultado)
    except: return texto

def melhorar_visual_topicos(texto_html):
    linhas = re.split(r'(<br>|\n)', texto_html)
    novo_texto = []
    for linha in linhas:
        if re.search(r'^\s*[-•*]\s+', re.sub(r'<[^>]+>', '', linha).strip()):
            linha_limpa = re.sub(r'^\s*[-•*]\s+', '', linha)
            novo_texto.append(f'<div class="topico-item">• {linha_limpa}</div>')
        else:
            novo_texto.append(linha)
    return "".join(novo_texto)

def destacar_datas(texto):
    padrao = r'(Esta\s+bula\s+foi\s+(?:atualizada\s+conforme\s+Bula\s+Padr[oã]o\s+)?aprovada\s+pela\s+Anvisa\s+em\s*)(\d{2}/\d{2}/\d{4}|\d{2}/\d{4})'
    def replacer(match):
        return f'{match.group(1)}<span class="highlight-blue">{match.group(2)}</span>'
    return re.sub(padrao, replacer, texto, count=0, flags=re.IGNORECASE | re.DOTALL)

# ----------------- 7. DIFF E REGRAS DE DECISÃO -----------------
def diff_palavra_a_palavra(texto_ref, texto_novo):
    ref_sem_tags = re.sub(r'<[^>]+>', '', texto_ref)
    novo_sem_tags = re.sub(r'<[^>]+>', '', texto_novo)
    ref_sem_tags = re.sub(r'[-–—]', ' ', ref_sem_tags)
    novo_sem_tags = re.sub(r'[-–—]', ' ', novo_sem_tags)
    palavras_ref = ref_sem_tags.split()
    palavras_novo = novo_sem_tags.split()
    matcher = difflib.SequenceMatcher(None, palavras_ref, palavras_novo)
    html_ref_list = []
    html_novo_list = []
    tem_diff = False
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            texto = " ".join(palavras_ref[i1:i2]); html_ref_list.append(texto); html_novo_list.append(texto)
        elif tag == 'replace':
            html_ref_list.append(f'<span class="highlight-yellow">{" ".join(palavras_ref[i1:i2])}</span>')
            html_novo_list.append(f'<span class="highlight-yellow">{" ".join(palavras_novo[j1:j2])}</span>')
            tem_diff = True
        elif tag == 'delete':
            html_ref_list.append(f'<span class="highlight-yellow">{" ".join(palavras_ref[i1:i2])}</span>')
            tem_diff = True
        elif tag == 'insert':
            html_novo_list.append(f'<span class="highlight-yellow">{" ".join(palavras_novo[j1:j2])}</span>')
            tem_diff = True
    return " ".join(html_ref_list), " ".join(html_novo_list), tem_diff

# --- SUA NOVA FUNÇÃO INTEGRADA AQUI ---
def gerar_diff_html(texto_ref, texto_novo):
    if not texto_ref: texto_ref = ""
    if not texto_novo: texto_novo = ""
    
    # 1. CHECAGEM NUCLEAR: Se o conteúdo alfanumérico for igual, ignora formatação
    if normalizacao_nuclear(texto_ref) == normalizacao_nuclear(texto_novo):
        html_novo = melhorar_visual_topicos(texto_novo.replace('\n', '<br>'))
        return texto_ref.replace('\n', '<br>'), html_novo, False

    # 2. Se falhar na nuclear, faz o diff detalhado
    ref_limpo = re.sub(r'<[^>]+>', '', texto_ref)
    novo_limpo = re.sub(r'<[^>]+>', '', texto_novo)
    
    r_html, n_html, diff_bool = diff_palavra_a_palavra(ref_limpo, novo_limpo)
    
    n_html_final = melhorar_visual_topicos(n_html)
    r_html_final = r_html.replace('\n', '<br>')
    
    return r_html_final, n_html_final, diff_bool

# ----------------- 8. EXTRAÇÃO DE TEXTO LOCAL -----------------
def extract_text_from_file(uploaded_file):
    try:
        text = ""
        name = getattr(uploaded_file, "name", "uploaded")
        if name.lower().endswith('.pdf'):
            uploaded_file.seek(0)
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            for page in doc:
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                for block in blocks:
                    if block.get("type") != 0: continue
                    block_text = ""
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            content = span.get("text", "")
                            if not content.strip(): line_text += content; continue
                            flags = span.get("flags", 0)
                            font_name = span.get("font", "").lower()
                            is_bold = ((flags & 16) or "bold" in font_name or "black" in font_name or "heavy" in font_name or "semibold" in font_name)
                            is_italic = ((flags & 2) or "italic" in font_name or "oblique" in font_name)
                            formatted_text = content
                            if is_bold and is_italic: formatted_text = f"<b><i>{content}</i></b>"
                            elif is_bold: formatted_text = f"<b>{content}</b>"
                            elif is_italic: formatted_text = f"<i>{content}</i>"
                            line_text += formatted_text
                        block_text += line_text.strip() + " "
                    text += block_text.strip() + "\n\n"
            doc.close()
        elif name.lower().endswith('.docx'):
            try:
                uploaded_file.seek(0); doc = docx.Document(uploaded_file)
            except:
                uploaded_file.seek(0); doc = docx.Document(BytesIO(uploaded_file.read()))
            for para in doc.paragraphs:
                para_text = ""
                for run in para.runs:
                    content = run.text
                    if not content: continue
                    is_bold = run.bold is True
                    is_italic = run.italic is True
                    formatted_text = content
                    if is_bold and is_italic: formatted_text = f"<b><i>{content}</i></b>"
                    elif is_bold: formatted_text = f"<b>{content}</b>"
                    elif is_italic: formatted_text = f"<i>{content}</i>"
                    para_text += formatted_text
                text += para_text + "\n\n"
        return clean_metadata_and_footers(text.strip())
    except Exception as e:
        st.error(f"Erro ao extrair texto do arquivo {getattr(uploaded_file, 'name', '')}: {str(e)}")
        return ""

# ----------------- 9. FUNÇÃO OCR (GEMINI) -----------------
def ocr_via_gemini(uploaded_file, keys_validas):
    uploaded_file.seek(0)
    bytes_data = uploaded_file.read()
    mime_type = "application/pdf"
    if getattr(uploaded_file, "name", "").lower().endswith(".docx"): return "" 
    
    prompt_ocr = """EXTRAIA TODO O TEXTO DESTE ARQUIVO.
    Mantenha formatação <b> e <i> onde apropriado.
    Não adicione comentários, apenas devolva o texto extraído."""
    
    for key in keys_validas:
        genai.configure(api_key=key)
        model = genai.GenerativeModel("gemini-2.5-flash")
        try:
            response = model.generate_content([{'mime_type': mime_type, 'data': bytes_data}, prompt_ocr])
            return clean_metadata_and_footers(response.text)
        except Exception: continue
    return ""

# ----------------- 10. CLEAN JSON HELPER -----------------
def clean_json_string(text):
    """Remove blocos de código Markdown (```json ... ```) e limpa espaços."""
    if not text: return ""
    match = re.search(r"```json\s*(.*?)\s*```", text, re.DOTALL)
    if match: return match.group(1).strip()
    match_generic = re.search(r"```\s*(.*?)\s*```", text, re.DOTALL)
    if match_generic: return match_generic.group(1).strip()
    return text.strip()

# ----------------- 11. UI PRINCIPAL E FLUXO -----------------
st.title("💊 Gráfica x Arte")
tipo_bula = st.radio("Escolha o Tipo de Bula:", ("Paciente",), horizontal=True)

c1, c2 = st.columns(2)
f1 = c1.file_uploader("📜 Gráfica", type=["pdf", "docx"], key="f1")
f2 = c2.file_uploader("📜 Arte Vigente", type=["pdf", "docx"], key="f2")

if st.button("🚀 Processar Conferência", key="process_button"):
    st.info("Iniciando processamento...")
    if not f1 or not f2:
        st.warning("Por favor, envie ambos os arquivos antes de processar."); st.stop()
    
    secret_names = ["GEMINI_API_KEY", "GEMINI_API_KEY1", "GEMINI_API_KEY2", "GEMINI_API_KEY3", "GEMNI_API_KEY1", "GEMNI_API_KEY2", "GEMNI_API_KEY3"]
    keys_raw = [st.secrets.get(n) for n in secret_names]
    keys_validas = [k for k in keys_raw if k]

    if not keys_validas:
        st.error("Erro crítico: nenhuma API key encontrada nos secrets."); st.stop()

    if f1 and f2:
        secoes_alvo = SECOES_PACIENTE
        with st.spinner("Lendo arquivos e analisando conteúdo..."):
            try:
                f1.seek(0); f2.seek(0)
                t_anvisa = extract_text_from_file(f1)
                t_mkt = extract_text_from_file(f2)
                
                # OCR Trigger
                if len(t_anvisa) < 1000:
                    st.info(f"Gráfica: Texto local insuficiente ({len(t_anvisa)} chars). Acionando OCR...")
                    ocr_res = ocr_via_gemini(f1, keys_validas)
                    if len(ocr_res) > len(t_anvisa): t_anvisa = ocr_res; st.success("OCR Gráfica OK.")
                    else: st.warning("OCR Gráfica não melhorou o resultado.")
                if len(t_mkt) < 1000:
                    st.info(f"Arte: Texto local insuficiente ({len(t_mkt)} chars). Acionando OCR...")
                    ocr_res2 = ocr_via_gemini(f2, keys_validas)
                    if len(ocr_res2) > len(t_mkt): t_mkt = ocr_res2; st.success("OCR Arte OK.")
                    else: st.warning("OCR Arte não melhorou o resultado.")

                st.info(f"Tamanho final Gráfica: {len(t_anvisa)} caracteres")
                st.info(f"Tamanho final Arte: {len(t_mkt)} caracteres")

                input_grafica = t_anvisa[:150000]
                input_arte = t_mkt[:150000]

            except Exception as e:
                st.exception(f"Falha ao processar arquivos: {e}"); st.stop()

            if len(t_anvisa) < 20 or len(t_mkt) < 20:
                st.error("Conteúdo insuficiente para análise."); st.stop()

            # PROMPT CORRIGIDO (JSON PURO)
            prompt = f"""
            Você é um Extrator de Dados Farmacêuticos Rigoroso.

            INPUT TEXTO 1 (GRÁFICA): {input_grafica}
            INPUT TEXTO 2 (ARTE VIGENTE): {input_arte}

            TAREFA: Extrair o texto de cada seção listada abaixo para ambos os inputs.

            REGRAS CRÍTICAS:
            1. Responda APENAS em JSON válido.
            2. NÃO use formatação Markdown (sem ```json).
            3. Mantenha <b> e <i>. Converta markdown **bold** para <b>bold</b>.
            4. Se o texto for muito longo, NÃO TRUNQUE a resposta no meio de uma string JSON.
            5. LISTA DE SEÇÕES ESPERADAS: {secoes_alvo}

            SAÍDA JSON ESPERADA:
            {{
             "secoes": [
                 {{
                     "titulo": "NOME DA SEÇÃO",
                     "texto_anvisa": "conteúdo completo da Gráfica",
                     "texto_mkt": "conteúdo completo da Arte Vigente"
                 }}
             ]
            }}
            """

            response = None
            sucesso = False
            log_erros = []
            extracted_date_ref = "-"
            extracted_date_mkt = "-"

            for idx_key, key in enumerate(keys_validas):
                if sucesso: break
                try:
                    genai.configure(api_key=key)
                except Exception as e:
                    log_erros.append(f"Key {idx_key+1} config error: {str(e)}"); continue

                for modelo in MODELOS_PARA_TENTAR:
                    try:
                        st.info(f"Analisando com modelo: {modelo} (Key {idx_key+1})...")
                        model = genai.GenerativeModel(
                            modelo,
                            generation_config={
                                "response_mime_type": "application/json", 
                                "temperature": 0.1,
                                "max_output_tokens": 65536 
                            }
                        )
                        response = model.generate_content(prompt)
                        sucesso = True
                        break
                    except Exception as e:
                        log_erros.append(f"Key {idx_key+1} | {modelo}: {str(e)}")
                        time.sleep(1); continue

            if not sucesso:
                st.error("❌ Falha na API Gemini."); st.code("\n".join(log_erros)); st.stop()

            try:
                resp_text = getattr(response, "text", None) or str(response)
                resp_text_limpo = clean_json_string(resp_text)
                
                try:
                    resultado = json.loads(resp_text_limpo)
                except json.JSONDecodeError as e:
                    try:
                        end_idx = resp_text_limpo.rfind('}')
                        if end_idx != -1:
                            resultado = json.loads(resp_text_limpo[:end_idx+1])
                        else:
                            raise e
                    except:
                        st.error(f"Erro JSON Inválido: {e}")
                        st.text_area("Dump Resposta IA (Parcial)", resp_text_limpo[:1000])
                        st.stop()

            except Exception as e:
                st.exception(f"Erro ao processar resposta: {e}"); st.stop()

            try:
                dados_secoes = resultado.get("secoes", [])
                secoes_finais = []
                divs_count = 0
                frase_padrao = r'Esta\s+bula\s+foi\s+atualizada\s+conforme\s+Bula\s+Padr[oã]o\s+aprovada\s+pela\s+Anvisa\s+em\s*(\d{2}/\d{2}/\d{4}|\d{2}/\d{4})'

                for item in dados_secoes:
                    titulo = item.get('titulo', '').strip()
                    txt_ref = item.get('texto_anvisa', '').strip()
                    txt_mkt = item.get('texto_mkt', '').strip()

                    # Fallback extração local
                    if not txt_ref:
                        tentativa = extract_section_from_raw(t_anvisa, titulo, secoes_alvo)
                        if tentativa: txt_ref = tentativa
                    if not txt_mkt:
                        tentativa2 = extract_section_from_raw(t_mkt, titulo, secoes_alvo)
                        if tentativa2: txt_mkt = tentativa2

                    txt_ref = clean_metadata_and_footers(txt_ref)
                    txt_mkt = clean_metadata_and_footers(txt_mkt)
                    txt_ref = convert_markdown_bold_to_html(txt_ref)
                    txt_mkt = convert_markdown_bold_to_html(txt_mkt)

                    if "DIZERES LEGAIS" in titulo.upper():
                        m_ref = re.search(frase_padrao, txt_ref, flags=re.IGNORECASE)
                        if m_ref: extracted_date_ref = m_ref.group(1)
                        m_mkt = re.search(frase_padrao, txt_mkt, flags=re.IGNORECASE)
                        if m_mkt: extracted_date_mkt = m_mkt.group(1)

                        html_ref = destacar_datas(txt_ref).replace('\n', '<br>')
                        html_novo = destacar_datas(txt_mkt).replace('\n', '<br>')
                        html_ref = verificar_ortografia_inteligente(html_ref)
                        html_novo = verificar_ortografia_inteligente(html_novo)
                        html_ref = melhorar_visual_topicos(html_ref)
                        html_novo = melhorar_visual_topicos(html_novo)
                        status = "CONFORME"
                    else:
                        # AQUI ESTAVA A MUDANÇA PRINCIPAL NA CHAMADA
                        html_ref, html_novo, teve_diff = gerar_diff_html(txt_ref, txt_mkt)
                        status = "DIVERGENTE" if teve_diff else "CONFORME"
                        if teve_diff: divs_count += 1

                    secoes_finais.append({
                        "titulo": titulo,
                        "texto_anvisa": html_ref,
                        "texto_mkt": html_novo,
                        "status": status
                    })

                data_ref = extracted_date_ref if extracted_date_ref != "-" else "-"
                data_mkt = extracted_date_mkt if extracted_date_mkt != "-" else "-"

                st.markdown("### 📊 Resumo")
                c1, c2, c3 = st.columns(3)
                c1.metric("Data Gráfica", data_ref)
                c2.metric("Data Arte Vigente", data_mkt, delta="Igual" if data_ref == data_mkt and data_ref != "-" else "Diferente")
                c3.metric("Seções", len(secoes_finais))

                sub1, sub2 = st.columns(2)
                sub1.info(f"✅ Conformes: {len(secoes_finais) - divs_count}")
                if divs_count > 0:
                    sub2.warning(f"⚠️ Divergentes: {divs_count}")
                else:
                    sub2.success("✨ Divergências: 0")

                st.divider()

                for item in secoes_finais:
                    status = item['status']
                    titulo = item['titulo']
                    if "DIZERES LEGAIS" in titulo.upper():
                        icon, css, aberto = "⚖️", "border-info", False
                    elif status == "CONFORME":
                        icon, css, aberto = "✅", "border-ok", False
                    else:
                        icon, css, aberto = "⚠️", "border-warn", True

                    with st.expander(f"{icon} {titulo}", expanded=aberto):
                        ce, cd = st.columns(2)
                        with ce:
                            st.caption("Gráfica")
                            st.markdown(f'<div class="texto-box {css}">{item["texto_anvisa"]}</div>', unsafe_allow_html=True)
                        with cd:
                            st.caption("Arte Vigente")
                            st.markdown(f'<div class="texto-box {css}">{item["texto_mkt"]}</div>', unsafe_allow_html=True)

            except Exception as e:
                st.exception(f"Erro ao processar resultado final: {e}"); st.stop()
else:
    st.info("Aguardando arquivos.")
